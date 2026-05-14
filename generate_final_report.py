"""Final raporu oluşturur: BioPreserve-AI Haziran 2026"""
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Sayfa kenar boşlukları ──────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Stil yardımcıları ───────────────────────────────────────────────────────
def set_font(run, size=12, bold=False, color=None):
    run.font.name  = "Times New Roman"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p

def add_body(doc, text, indent=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=12)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, size=12)
    return p

def add_image_placeholder(doc, caption):
    """Görsel için gri yer tutucu kutu."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    border_run = p.add_run(
        f"\n\n[  GÖRSEL ALANINA EKLENECEK: {caption}  ]\n\n"
    )
    border_run.font.name  = "Courier New"
    border_run.font.size  = Pt(10)
    border_run.font.color.rgb = RGBColor(120, 120, 120)

    # Kutu çizgisi
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "12")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), "A0A0A0")
        pBdr.append(el)
    pPr.append(pBdr)
    return p

def simple_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=11, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=11)
    return table

# ═══════════════════════════════════════════════════════════════════════════
#  KAPAK
# ═══════════════════════════════════════════════════════════════════════════
for line in [
    ("T.C. İSTANBUL SABAHATTİN ZAİM ÜNİVERSİTESİ", 13, True),
    ("MÜHENDİSLİK VE DOĞA BİLİMLERİ FAKÜLTESİ", 13, True),
    ("", 12, False),
    ("BioPreserve-AI: Gıda Güvenliğinde Yapay Zeka Destekli\nDoğal Koruyucu Keşfi", 16, True),
    ("", 12, False),
    ("TASARIM PROJESİ – FİNAL RAPORU", 13, True),
    ("", 12, False),
    ("BETÜL PINARCI, SAİD TAHA DEMİR, BÜŞRA DEMİR, BEYZA BARTİN", 12, True),
    ("Gıda Mühendisliği | Bilgisayar Mühendisliği | Moleküler Biyoloji ve Genetik", 12, False),
    ("", 12, False),
    ("Danışman: Dr. Öğr. Üyesi Cem TURAN", 12, False),
    ("", 12, False),
    ("HAZİRAN 2026", 13, True),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line[0])
    set_font(run, size=line[1], bold=line[2])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  ÖNSÖZ
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "ÖNSÖZ", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_body(doc,
    "Bu çalışma, gıda endüstrisinde sürdürülebilir ve sağlıklı koruyucu sistemlerin "
    "geliştirilmesi amacıyla disiplinler arası bir yaklaşımla hazırlanmıştır. Nisan 2026'da "
    "sunulan araştırma dökümanının devamı niteliğindeki bu final raporu, sistemin "
    "tamamlanmış ve çalışır haldeki versiyonunu yansıtmaktadır. Proje sürecindeki "
    "desteklerinden dolayı danışmanımıza ve tüm ekip arkadaşlarımıza teşekkür ederiz."
)
doc.add_paragraph()
for name in ["Betül Pınarcı", "Sait Taha Demir", "Büşra Demir", "Beyza Bartin"]:
    p = doc.add_paragraph()
    run = p.add_run(name)
    set_font(run, size=12)
p = doc.add_paragraph()
run = p.add_run("HAZİRAN 2026")
set_font(run, size=12)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  KISALTMALAR
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "KISALTMALAR", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
for item in [
    "AMP  : Antimikrobiyal Peptit",
    "API  : Application Programming Interface",
    "DBAASP : Database of Antimicrobial Activity and Structure of Peptides",
    "DRAMP : Data Repository of Antimicrobial Peptides",
    "ETL  : Extract, Transform, Load",
    "MBG  : Moleküler Biyoloji ve Genetik",
    "MIC  : Minimum İnhibitör Konsantrasyonu",
    "ML   : Machine Learning (Makine Öğrenmesi)",
    "RF   : Random Forest",
    "UI   : User Interface (Kullanıcı Arayüzü)",
]:
    add_body(doc, item)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  TABLO & ŞEKİL LİSTELERİ
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "TABLO LİSTESİ", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
for t in [
    "Tablo 3.1 : Biyolojik Filtreleme Parametreleri ve Uygulanan Eşik Değerleri",
    "Tablo 3.2 : Gıda Matrisi Koşulları",
    "Tablo 3.3 : Hedef Patojen Profilleri",
    "Tablo 4.1 : Eğitim Veri Seti İstatistikleri",
    "Tablo 5.1 : Çok Aşamalı Filtreleme Sonuçları — 1.847 Test Dizisi Üzerinde Gerçek Değerler",
    "Tablo 5.2 : Model Sınıflandırma Metrikleri — 1.847 Test Dizisi",
    "Tablo 5.3 : Proje Çalışma Takvimi ve Gerçekleşme Durumu",
]:
    add_body(doc, t)
doc.add_paragraph()
add_heading(doc, "ŞEKİL LİSTESİ", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
for s in [
    "Şekil 3.1 : Sistem Genel Akış Diyagramı (Pipeline)",
    "Şekil 3.2 : Çok Aşamalı Biyolojik Filtreleme Şeması",
    "Şekil 4.1 : Yazılım Mimari Diyagramı",
    "Şekil 4.2 : Web Arayüzü – Analiz Seçim Ekranı",
    "Şekil 4.3 : Web Arayüzü – Sonuç Tablosu",
    "Şekil 4.4 : Web Arayüzü – Yeni Gıda/Patojen Ekleme Formu",
    "Şekil 5.1 : Random Forest Özellik Önem Grafiği",
    "Şekil 5.2 : Model Karışıklık Matrisi",
    "Şekil 5.3 : Onaylanan ve Elenen Adayların Dağılımı",
]:
    add_body(doc, s)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "ABSTRACT", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_body(doc,
    "This report presents the completed BioPreserve-AI system — a machine learning-assisted "
    "decision-support tool for discovering natural antimicrobial peptides (AMPs) as alternatives "
    "to synthetic food preservatives. The system integrates food engineering, molecular biology, "
    "and computer science into a fully operational, web-based pipeline. "
    "A Random Forest classifier was trained on 7,387 sequences (80% of 9,234 total) from the DBAASP "
    "and DRAMP databases, with 1,847 sequences (20%) held out as an independent test set. "
    "The classifier achieved 69.6% accuracy on this test set (precision 64.9%, recall 62.0% for "
    "AMP class). When the full pipeline was applied to the 1,847 test sequences, multi-stage "
    "biological rule filtering eliminated 1,528 sequences (82.7%), the ML stage eliminated a "
    "further 113 (6.1%), and 206 sequences (11.2%) were approved as candidate preservatives. "
    "A FastAPI backend and a React (Vite) frontend provide an interactive interface that supports "
    "system peptide pool analysis, CSV upload, food matrix selection, and target pathogen selection. "
    "Food engineers can extend the system by adding new food matrices and pathogen profiles directly "
    "through the UI, without modifying source code. The delivered prototype reduces the "
    "pre-laboratory candidate pool to a manageable shortlist within seconds, offering a "
    "cost-effective and sustainable path toward clean-label food preservation."
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  ÖZET
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "ÖZET", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_body(doc,
    "Bu proje, gıda endüstrisinde sentetik koruyuculara (sodyum nitrit, sorbik asit vb.) "
    "alternatif olabilecek doğal antimikrobiyal peptitleri (AMP) hızlı ve düşük maliyetle "
    "belirlemek amacıyla geliştirilen bir karar destek sistemi olan BioPreserve-AI'ın "
    "tamamlanmış halini sunmaktadır."
)
add_body(doc,
    "Sistem; DBAASP ve DRAMP veritabanlarından derlenen toplam 9.234 peptit dizisine "
    "dayanmaktadır. Verinin %80'i (7.387 dizi) model eğitimi, %20'si (1.847 dizi) ise "
    "bağımsız test değerlendirmesi için ayrılmıştır. Model, peptit uzunluğu, net yük, "
    "hidrofobisite oranı ve anahtar amino asit oranı olmak üzere dört biyofiziksel "
    "özelliği kullanmakta; test kümesinde %69.6 doğruluk sergilemektedir (AMP sınıfı "
    "için kesinlik %64.9, duyarlılık %62.0)."
)
add_body(doc,
    "Çok aşamalı biyolojik filtre, ML'den önce aday havuzunu dramatik biçimde daraltmaktadır. "
    "1.847 test dizisi üzerinde yürütülen tam pipeline değerlendirmesinde biyolojik filtreler "
    "1.528 diziyi (%82.7) eleyerek yalnızca 319 diziyi ML aşamasına iletti; ML modeli "
    "bu 319 diziden 113'ünü daha eleyerek 206 adayı (%11.2) 'Onaylandı' olarak işaretledi."
)
add_body(doc,
    "Kullanıcı arayüzü, sistem peptit havuzunu veya kullanıcının yüklediği CSV dosyasını "
    "analiz edebilmekte; analiz odağını (genel, gıda odaklı, patojen odaklı veya ikisi "
    "birlikte) seçebilmektedir. Gıda mühendisleri, kaynak kodu değiştirmeksizin arayüzden "
    "yeni gıda matrisi ve patojen profili ekleyebilmektedir. Eklenen veriler kalıcı olarak "
    "sisteme kaydedilmekte ve anında analizde kullanılabilmektedir."
)
add_body(doc,
    "Sistem; Listeria monocytogenes, Salmonella ve E. coli başta olmak üzere gıda güvenliği "
    "açısından kritik patojenleri, Peynir, Yoğurt, Taze Et, Fermente Et, Meyve Suyu ve "
    "Bitkisel Protein Bazlı Ürün gibi farklı gıda matrislerine göre değerlendirmektedir."
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  1. GİRİŞ
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. GİRİŞ")
add_body(doc,
    "Gıda endüstrisinde ürünlerin raf ömrünü uzatmak için uzun zamandır sodyum nitrit, "
    "sorbik asit gibi sentetik koruyucular kullanılmaktadır. Ancak son yıllarda bu katkı "
    "maddelerine karşı tüketici tarafında ciddi bir mesafe oluşmaktadır. 'Temiz etiket' "
    "(clean label) yaklaşımının yaygınlaşmasıyla birlikte insanlar, ambalajlardaki "
    "kimyasal isimlere karşı daha temkinli davranmaya başlamıştır."
)
add_body(doc,
    "Bu durum yalnızca bir algı meselesi değildir. Bazı araştırmalar bu tür katkıların "
    "uzun vadede bağırsak florasını olumsuz etkileyebileceğini göstermektedir. Bunun yanı "
    "sıra Listeria monocytogenes, Salmonella ve E. coli gibi patojenler, maruz kaldıkları "
    "koruyuculara zamanla direnç geliştirebildiğinden endüstri sürekli yeni çözümler "
    "aramak zorunda kalmaktadır."
)
add_body(doc,
    "Antimikrobiyal peptitler (AMP'ler), doğada birçok canlının savunma sisteminde bulunan "
    "kısa protein dizileridir. Bakterilerin hücre zarını fiziksel olarak hedef aldıklarından "
    "direnç gelişimi çok daha yavaştır. Ancak on binlerce peptit arasından doğru adayı "
    "bulmak, geleneksel laboratuvar yöntemleriyle hem zaman hem de maliyet açısından "
    "oldukça zahmetlidir."
)
add_body(doc,
    "BioPreserve-AI, bu ön-eleme sürecini in silico (bilgisayar ortamında) gerçekleştiren "
    "bir karar destek sistemidir. Bilgisayar mühendisliği (algoritma ve yazılım altyapısı), "
    "gıda mühendisliği (hedef gıda ve patojen seçimi) ve moleküler biyoloji & genetik "
    "(biyolojik eşik değerleri) disiplinlerini birleştirerek geliştirilmiştir. Bu final "
    "raporu, Nisan 2026'da sunulan araştırma dökümanında planlanan sistemin tüm bileşenleriyle "
    "tamamlanmış halini belgelemektedir."
)

# ═══════════════════════════════════════════════════════════════════════════
#  2. LİTERATÜR TARAMASI
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. LİTERATÜR TARAMASI")
add_body(doc,
    "AMP'lerin gıdada kullanımına yönelik bilimsel ilgi son birkaç yılda belirgin biçimde "
    "artmıştır. Xie ve ark. (2025), AMP tabanlı biyomalzemelerin gıda güvenliği alanında "
    "nasıl kullanılabileceğini incelemiş ve bu peptitlerin ambalaj sistemleriyle birlikte "
    "uygulanabileceğini ortaya koymuştur. Wang ve ark. (2025) ise yapay zeka ve phage "
    "display tekniklerini kullanarak biyoaktif peptit dizilerinin daha hızlı taranıp "
    "optimize edilebileceğini göstermiştir."
)
add_body(doc,
    "Deneysel kanıtlar açısından Rai ve ark. (2016), AMP'lerin bakteri zarını hedef alarak "
    "hızlı antimikrobiyal etki gösterdiğini ve geleneksel koruyuculara ciddi bir alternatif "
    "oluşturduğunu belirlemiştir. Huan ve ark. (2020) ise AMP'lerin sınıflandırılmasını, "
    "tasarım yöntemlerini ve etki mekanizmalarını incelemiş; biyoinformatik araçların bu "
    "alandaki rolünün kritik olduğunu vurgulamıştır."
)
add_body(doc,
    "Hesaplamalı yaklaşımlar açısından da önemli çalışmalar mevcuttur. Hu ve ark. (2025), "
    "Lactococcus lactis'ten elde edilen bir peptidin E. coli ve Salmonella gibi Gram-negatif "
    "patojenlere karşı yüksek sıcaklık ve tuz ortamında bile stabil kaldığını tespit etmiştir. "
    "Dai ve ark. (2026) da Lactiplantibacillus plantarum'un geniş spektrumlu patojenleri "
    "farklı mekanizmalarla inhibe ettiğini göstermiştir."
)
add_body(doc,
    "Zhang ve ark. (2024) AMP'lerin üretim maliyeti ve stabilitesine ilişkin zorlukların "
    "hâlâ sürdüğüne dikkat çekerken, Oraç (2024) gıda işleme atıklarından enzimatik yolla "
    "elde edilen biyoaktif peptitlerin hem ekonomik hem çevresel avantaj sunduğunu "
    "savunmuştur. Tüm bu çalışmalar, laboratuvar öncesi sistematik bir tarama aracının "
    "gerekliliğini ortaya koymakta; projemiz de tam bu boşluğu doldurmaktadır."
)

# ═══════════════════════════════════════════════════════════════════════════
#  3. YÖNTEM VE MODEL
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. YÖNTEM VE MODEL")

add_heading(doc, "3.1 Problemin Tanımlanması", size=13)
add_body(doc,
    "Gıda endüstrisi için uygun doğal antimikrobiyal peptitleri bulmak hem zaman alıcı "
    "hem de maliyetlidir. BioPreserve-AI, bu süreci hızlandırmak amacıyla çok aşamalı "
    "biyolojik filtreleme ile makine öğrenmesini birleştiren hibrit bir ön-eleme "
    "sistemi olarak tasarlanmıştır. Laboratuvar aşamasını ortadan kaldırmak değil, "
    "oraya gitmeden önce en umut verici adayları belirlemek temel hedeftir."
)

add_heading(doc, "3.2 Biyolojik Filtreleme Parametreleri", size=13)
add_body(doc,
    "Sistem, her peptiti aşağıdaki biyolojik kurallara göre değerlendirmektedir. "
    "Tüm eşik değerleri gıda mühendisliği ve MBG ekibi tarafından literatür "
    "taramasına dayalı olarak belirlenmiştir:"
)
doc.add_paragraph()
simple_table(doc,
    ["Parametre", "Uygulanan Eşik", "Gerekçe"],
    [
        ["Peptit uzunluğu", "15–30 amino asit", "Kısa diziler membrana nüfuz edemez; uzun diziler toksik risk taşır"],
        ["Net yük", "+3 ile +9", "Pozitif yük, negatif yüklü bakteri zarına elektrostatik bağlanmayı sağlar"],
        ["Hidrofobisite oranı", "%40–65", "Zar etkileşimi için gerekli; çok yüksek değerler sitotoksisite riski doğurur"],
        ["Anahtar AA oranı (K,R,L,I,F)", "≥ %25", "Antimikrobiyal aktiviteyle en çok ilişkilendirilen amino asitler"],
    ]
)
doc.add_paragraph()
add_body(doc, "Tablo 3.1: Biyolojik Filtreleme Parametreleri ve Uygulanan Eşik Değerleri")

add_heading(doc, "3.3 Gıda Matrisi ve Patojen Profilleri", size=13)
add_body(doc,
    "Her gıda matrisi için pH, tuz oranı, sıcaklık ve yağ içeriği parametreleri "
    "tanımlanmıştır. Sistem bu değerlere göre peptitin ilgili gıdayla uyumlu olup "
    "olmadığını değerlendirir:"
)
doc.add_paragraph()
simple_table(doc,
    ["Gıda Ürünü", "pH", "Tuz Oranı (%)", "Sıcaklık (°C)", "Yağ İçeriği"],
    [
        ["Peynir",                    "5.0", "2.5", "4",  "Yüksek"],
        ["Yoğurt",                   "4.5", "0.5", "4",  "Orta"],
        ["Taze Et",                  "5.6", "1.2", "4",  "Yüksek"],
        ["Fermente Et (Sucuk)",       "4.8", "3.0", "8",  "Yüksek"],
        ["Meyve Suyu",               "3.8", "0.1", "10", "Düşük"],
        ["Bitkisel Protein Bazlı Ürün","6.2","0.8", "5",  "Orta"],
    ]
)
doc.add_paragraph()
add_body(doc, "Tablo 3.2: Gıda Matrisi Koşulları")
doc.add_paragraph()
add_body(doc, "Hedef patojen profilleri ise aşağıdaki şekilde tanımlanmıştır:")
doc.add_paragraph()
simple_table(doc,
    ["Patojen", "Min. Net Yük", "Hidrofobisite Aralığı", "Tip"],
    [
        ["Listeria monocytogenes", "≥ 3", "%40–65", "Gram-pozitif"],
        ["Salmonella",            "≥ 4", "%35–60", "Gram-negatif"],
        ["E. coli",               "≥ 4", "%35–60", "Gram-negatif"],
    ]
)
doc.add_paragraph()
add_body(doc, "Tablo 3.3: Hedef Patojen Profilleri")

add_heading(doc, "3.4 Makine Öğrenmesi Modeli", size=13)
add_body(doc,
    "Biyolojik filtreleri geçen adaylar, bir Random Forest sınıflandırıcısına iletilmektedir. "
    "Modelin temel özellikleri şu şekildedir:"
)
for b in [
    "Algoritma: Random Forest (120 ağaç, maksimum derinlik 7, minimum yaprak örneği 1)",
    "Özellikler: peptit uzunluğu, net yük, hidrofobisite oranı, anahtar amino asit oranı",
    "Toplam veri: 9.234 dizi (DBAASP ve DRAMP veritabanları)",
    "Eğitim seti: 7.387 dizi (%80, stratified ayırım, random_state=42)",
    "Test seti: 1.847 dizi (%20) — pipeline varsayılan girdi olarak da kullanılır",
    "Test doğruluğu: %69.6 — AMP sınıfı kesinlik %64.9 / duyarlılık %62.0",
    "Çıktı: 0–1 arasında AMP olasılık skoru (predict_proba[1])",
]:
    add_bullet(doc, b)
add_body(doc,
    "Model, yorumlanabilirlik ön planda tutularak tasarlanmıştır. Derin öğrenme modellerinin "
    "aksine Random Forest, hangi özelliklerin kararı etkilediğini açıkça ortaya koyabilmekte "
    "ve gıda mühendislerine anlamlı geri bildirim sağlamaktadır."
)

add_heading(doc, "3.5 Genel Sistem Akışı", size=13)
add_body(doc,
    "Sistemin uçtan uca akışı aşağıdaki şemada gösterilmektedir:"
)
add_image_placeholder(doc, "Şekil 3.1 – Sistem Genel Akış Diyagramı (Pipeline)")
add_body(doc,
    "Her aşamada elenen adaylar nedenleriyle birlikte raporlanmakta; yalnızca tüm "
    "aşamaları geçen peptitler 'Onaylandı' statüsüyle kullanıcıya sunulmaktadır."
)
add_image_placeholder(doc, "Şekil 3.2 – Çok Aşamalı Biyolojik Filtreleme Şeması")

# ═══════════════════════════════════════════════════════════════════════════
#  4. MODELİN DİJİTAL UYGULAMASI
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. MODELİN DİJİTAL UYGULAMASI")

add_heading(doc, "4.1 Yazılım Mimarisi", size=13)
add_body(doc,
    "BioPreserve-AI tek servis olarak dağıtılacak şekilde tasarlanmıştır. "
    "Sistem bileşenleri şu şekilde özetlenebilir:"
)
for b in [
    "Backend: Python 3.12 + FastAPI — REST API ve statik dosya sunumu",
    "Frontend: React 18 (Vite) — etkileşimli tek sayfa uygulaması",
    "ML: Scikit-learn RandomForestClassifier, joblib ile serileştirilmiş model",
    "Veri: JSON tabanlı kural dosyaları (biological_rules.json, food_context.json)",
    "Dağıtım: Render üzerinde tek servis (frontend build edilmiş haliyle backend'den sunulur)",
]:
    add_bullet(doc, b)
add_image_placeholder(doc, "Şekil 4.1 – Yazılım Mimari Diyagramı")

add_heading(doc, "4.2 Eğitim Veri Seti İstatistikleri", size=13)
add_body(doc,
    "Modelin eğitildiği veri seti DBAASP ve DRAMP veritabanlarından derlenerek "
    "temizlenmiş ve %80/%20 oranında eğitim/test olarak ayrılmıştır:"
)
doc.add_paragraph()
simple_table(doc,
    ["Parametre", "Değer"],
    [
        ["Toplam peptit dizisi",          "9.234"],
        ["Eğitim seti (%80)",             "7.387 dizi"],
        ["Test seti (%20)",               "1.847 dizi"],
        ["Test — AMP (label=1)",          "785 dizi (%42.5)"],
        ["Test — AMP değil (label=0)",    "1.062 dizi (%57.5)"],
        ["Kullanılan özellik sayısı",     "4 (biyofiziksel)"],
        ["Test doğruluğu",                "%69.6"],
        ["AMP sınıfı F1-skoru",           "0.635"],
    ]
)
doc.add_paragraph()
add_body(doc, "Tablo 4.1: Eğitim Veri Seti İstatistikleri")

add_heading(doc, "4.4 Pipeline Modülleri", size=13)
add_body(doc,
    "Yazılım üç ana modülden oluşmaktadır:"
)
add_body(doc,
    "features/feature_extraction.py — Verilen amino asit dizisinden dört biyofiziksel "
    "özelliği hesaplar: uzunluk, net yük (K/R/H için +1, D/E için -1), hidrofobisite oranı "
    "(A,V,I,L,M,F,W,Y), anahtar amino asit oranı (K,R,L,I,F). Geçersiz amino asit "
    "içeren diziler hata fırlatarak elenir."
)
add_body(doc,
    "pipeline.py — Biyolojik kural kontrolü, patojen uyumluluğu, gıda matrisi "
    "uyumluluğu ve ML tahminini sırayla uygular. Her peptit için kural notu ve gıda notu "
    "üretir; nihai durum 'Onaylandı' veya 'Elendi' olarak raporlanır."
)
add_body(doc,
    "backend/main.py — FastAPI uygulaması. Sistem peptit havuzu (GET /api/run-pipeline) "
    "ve CSV yükleme (POST /api/run-pipeline) endpoint'lerinin yanı sıra yeni gıda matrisi "
    "(POST /api/food-matrices) ve patojen (POST /api/pathogens) ekleme endpoint'lerini sunar."
)

add_heading(doc, "4.5 Web Arayüzü", size=13)
add_body(doc,
    "Kullanıcı arayüzü dört temel işlevi desteklemektedir: veri kaynağı seçimi "
    "(sistem havuzu veya CSV yükleme), analiz odağı seçimi (genel / gıda odaklı / "
    "patojen odaklı / ikisi birlikte), sonuç tablosunda filtreleme ve arama, "
    "CSV raporu indirme."
)
add_image_placeholder(doc, "Şekil 4.2 – Web Arayüzü: Analiz Seçim Ekranı")
add_image_placeholder(doc, "Şekil 4.3 – Web Arayüzü: Sonuç Tablosu")

add_heading(doc, "4.6 Gıda Mühendisi Uzantı Arayüzü", size=13)
add_body(doc,
    "Araştırma dökümanından bu yana eklenen en önemli yenilik, gıda mühendislerinin "
    "kaynak kod değiştirmeksizin sistemi genişletebilmesidir. Gıda odaklı veya birleşik "
    "modda '+ Yeni gıda' butonuna basıldığında bir form açılır; kullanıcı ürün adı, pH, "
    "tuz oranı, sıcaklık ve yağ içeriği bilgilerini girer. Kaydedilen veri food_context.json "
    "dosyasına kalıcı olarak yazılır ve dropdown listesine anında eklenir. Aynı mekanizma "
    "patojen ekleme için de geçerlidir."
)
add_image_placeholder(doc, "Şekil 4.4 – Yeni Gıda Matrisi / Patojen Ekleme Formu")

add_heading(doc, "4.7 Vaka Senaryosu", size=13)
add_body(doc,
    "Kullanıcının 'Taze Et'te Listeria'ya karşı etkili peptit aradığı bir senaryo "
    "düşünüldüğünde sistem şu adımları izler:"
)
for i, step in enumerate([
    "Sistem peptit havuzu (database_peptides.csv — 1.847 bağımsız test dizisi) yüklenir.",
    "Analiz odağı 'Gıda + patojen', hedef gıda 'Taze Et', hedef patojen 'Listeria monocytogenes' seçilir.",
    "Her dizi için özellik çıkarımı yapılır; biyolojik filtreler sırayla uygulanır.",
    "Filtreyi geçen adaylar ML modeline gönderilir; olasılık skoru hesaplanır.",
    "Listeria uyumluluğu ve Taze Et koşulları kontrol edilir.",
    "Tüm aşamaları geçen peptitler 'Onaylandı' olarak listelenir ve CSV olarak indirilebilir.",
], 1):
    add_body(doc, f"{i}. {step}", indent=0.5)

# ═══════════════════════════════════════════════════════════════════════════
#  5. ELDE EDİLEN SONUÇLARIN ANALİZİ VE TARTIŞMA
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. ELDE EDİLEN SONUÇLARIN ANALİZİ VE TARTIŞMA")

add_heading(doc, "5.1 Filtreleme Performansı", size=13)
add_body(doc,
    "1.847 test dizisi üzerinde gerçekleştirilen genel tarama (hedef gıda/patojen seçilmeksizin) "
    "şu sonuçları vermiştir:"
)
doc.add_paragraph()
simple_table(doc,
    ["Aşama", "Eleme Kriteri", "Elenen Dizi", "Kümülatif Oran"],
    [
        ["1–4. Biyolojik filtreler",    "Uzunluk, yük, hidrofobisite, anahtar AA", "1.528", "%82.7"],
        ["5. ML filtresi",              "predict_label ≠ 1",                        "113",   "%6.1"],
        ["6. Gıda/patojen filtresi",   "Genel taramada uygulanmaz",                "—",     "—"],
        ["Onaylanan aday",             "Tüm aşamaları geçti",                      "206",   "%11.2"],
    ]
)
doc.add_paragraph()
add_body(doc,
    "Tablo 5.1: Çok Aşamalı Filtreleme Sonuçları — 1.847 Test Dizisi Üzerinde Gerçek Değerler"
)
add_body(doc,
    "Biyolojik filtreler aday havuzunu ML aşamasına ulaşmadan %82.7 oranında daraltmış, "
    "ML ise kalan 319 adaydan 113'ünü (%35.4) daha eleyerek nihai listeyi 206 adaya indirmiştir. "
    "Ortalama ML olasılık skoru (yalnızca ML'e giren 319 dizi üzerinden) 0.54 olarak ölçülmüştür."
)
add_image_placeholder(doc, "Şekil 5.3 – Onaylanan ve Elenen Adayların Dağılımı (Pasta / Bar Grafiği)")

add_heading(doc, "5.2 Model Performansı", size=13)
add_body(doc,
    "Random Forest modeli 7.387 diziden oluşan eğitim seti ile eğitilmiş olup "
    "1.847 diziden oluşan bağımsız test kümesinde %69.6 doğruluk elde etmiştir. "
    "Modelin yalnızca dört biyofiziksel özellik kullandığı göz önünde bulundurulduğunda "
    "bu sonuç, seçilen özelliklerin AMP aktivitesiyle anlamlı ilişki içinde olduğunu "
    "doğrulamaktadır. Ayrıntılı sınıflandırma metrikleri aşağıda verilmiştir:"
)
doc.add_paragraph()
simple_table(doc,
    ["Sınıf", "Kesinlik (Precision)", "Duyarlılık (Recall)", "F1-Skoru", "Destek (Support)"],
    [
        ["AMP değil (label=0)", "%72.8", "%75.2", "0.740", "1.062"],
        ["AMP (label=1)",       "%64.9", "%62.0", "0.635",   "785"],
        ["Genel doğruluk",     "%69.6",     "—",     "—",   "1.847"],
        ["Makro ortalama",     "%68.9", "%68.6", "0.687",   "1.847"],
    ]
)
doc.add_paragraph()
add_body(doc, "Tablo 5.2: Model Sınıflandırma Metrikleri — 1.847 Test Dizisi")
add_image_placeholder(doc, "Şekil 5.1 – Random Forest Özellik Önem Grafiği (feature_importances_)")
add_image_placeholder(doc, "Şekil 5.2 – Model Karışıklık Matrisi")
add_body(doc,
    "Araştırma dökümanından bu yana ortalama ML skoru hesaplama yöntemi de iyileştirilmiştir: "
    "Ortalama artık yalnızca biyolojik filtreyi geçip gerçekten ML'e giren peptitler "
    "üzerinden hesaplanmakta, böylece elenen adayların (ml_probability = 0.0) ortalamayı "
    "yapay olarak düşürmesi engellenmektedir."
)

add_heading(doc, "5.3 Temel Gözlemler ve Tartışma", size=13)
for obs in [
    "Hibrit yaklaşım (kural + ML) işe yaramaktadır. Kural tabanlı filtreleme ML yükünü önemli ölçüde azaltırken, ML modeli kural filtrelerini geçen ancak gerçekte AMP olmayan adayları ek bir katmanla eleyebilmektedir.",
    "Hidrofobisite-aktivite dengesi kritiktir. Hidrofobisite arttıkça antimikrobiyal potansiyel artmakta, ancak sitotoksisite riski de yükselmektedir. Seçilen %40–65 aralığı bu dengeyi korumak üzere belirlenmiştir.",
    "Gıda koşulları adayları önemli ölçüde etkiler. Aynı peptit yüksek tuzlu bir ortamda (sucuk) başarısız olurken düşük tuzlu bir ortamda (meyve suyu) geçebilmektedir. Bu durum, gıda matrisi seçiminin analizin ayrılmaz bir parçası olduğunu doğrulamaktadır.",
    "Sistem genişletilebilirliği sağlanmıştır. Gıda mühendisleri kod yazmaksızın yeni gıda matrisi ve patojen ekleyebildiğinden sistem farklı araştırma senaryolarına kolayca uyarlanabilmektedir.",
]:
    add_bullet(doc, obs)

# ═══════════════════════════════════════════════════════════════════════════
#  6. PROJENİN GENEL DEĞERLENDİRMESİ
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. PROJENİN GENEL DEĞERLENDİRMESİ")

add_heading(doc, "6.1 Çalışma Takvimi Karşılaştırması", size=13)
doc.add_paragraph()
simple_table(doc,
    ["Dönem", "Planlanan Çalışmalar", "Gerçekleşme Durumu"],
    [
        ["Nisan 2026",
         "Veri çekme, temizleme, özellik çıkarımı, ML model eğitimi",
         "Tamamlandı — DBAASP & DRAMP verisi ETL'den geçirildi; model eğitildi ve kaydedildi"],
        ["Mayıs 2026",
         "Filtreleme motoru + ML entegrasyonu, FastAPI backend, React frontend",
         "Tamamlandı — Çok aşamalı pipeline, web arayüzü ve API tam işlevsel hale getirildi"],
        ["Haziran 2026 (İlk Hafta)",
         "Son testler, arayüz iyileştirmeleri, final raporu teslimi",
         "Tamamlandı — Türkçe notlar, ML skoru düzeltmesi, gıda/patojen ekleme arayüzü teslim edildi"],
    ]
)
doc.add_paragraph()
add_body(doc, "Tablo 5.3: Proje Çalışma Takvimi ve Gerçekleşme Durumu")

add_heading(doc, "6.2 Karşılaşılan Zorluklar", size=13)
for z in [
    "Veri format tutarsızlıkları: DBAASP ve DRAMP farklı yapılarda veri sunmaktadır; normalizasyon aşaması beklenenden uzun sürmüştür.",
    "Özellik mühendisliği sınırları: Dört biyofiziksel özellik yorumlanabilirlik açısından avantajlıdır, ancak biyolojik gerçekliği tam yansıtmamaktadır; bu durum ~%70 tavan doğruluğuna yol açmaktadır.",
    "Hidrofobisite-toksisite dengesi: Bu iki parametre arasındaki doğru eşiği belirlemek MBG ve gıda ekibiyle kapsamlı tartışmalar gerektirmiştir.",
    "Gıda matrisi uyum kuralları: Farklı gıdaların farklı pH, tuz ve yağ kombinasyonları, kural mantığının titizlikle kurgulanmasını zorunlu kılmıştır.",
]:
    add_bullet(doc, z)

add_heading(doc, "6.3 Yaygın Etki Değerlendirmesi", size=13)
add_body(doc,
    "Ekonomik açıdan sistem, büyük Ar-Ge bütçesi olmayan firmalar için laboratuvar "
    "denemesi sayısını ve maliyetini azaltmaktadır. Çevresel açıdan ise her başarısız "
    "deney kimyasal atık demek olduğundan ön-eleme mekanizması atık oluşumunu da "
    "dolaylı olarak azaltmaktadır. Sürdürülebilirlik boyutunda sentetik koruyuculara "
    "daha az bağımlılık, temiz etiket üretimi ve gıda israfının azalması gibi olumlu "
    "etkiler beklenmektedir. Sağlık boyutunda ise tüketicilerin maruz kaldığı sentetik "
    "katkı maddesi miktarının azalması bağırsak sağlığı ve genel toksisite riski "
    "açısından olumlu sonuçlar doğurabilecektir."
)

# ═══════════════════════════════════════════════════════════════════════════
#  7. SONUÇLAR VE ÖNERİLER
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "7. SONUÇLAR VE ÖNERİLER")

add_heading(doc, "7.1 Sonuçlar", size=13)
add_body(doc,
    "Bu çalışmada gıda güvenliği alanında sentetik koruyuculara alternatif doğal "
    "antimikrobiyal peptitleri hızla belirleyen, uçtan uca çalışan bir karar destek "
    "sistemi geliştirilmiş ve teslim edilmiştir. Sistem;"
)
for m in [
    "9.234 peptit içeren veri setinin %80'i (7.387 dizi) üzerinde eğitilmiş, test setinde %69.6 doğruluğa ulaşan bir Random Forest modeli,",
    "Biyolojik kurallara dayalı çok aşamalı bir filtreleme motoru,",
    "FastAPI tabanlı REST API ve React tabanlı etkileşimli web arayüzü,",
    "Gıda mühendislerinin kod yazmaksızın yeni gıda matrisi ve patojen ekleyebildiği bir genişletme mekanizması",
]:
    add_bullet(doc, m)
add_body(doc,
    "bileşenlerini başarıyla bir araya getirmektedir. Araştırma dökümanında planlanan "
    "tüm hedefler gerçekleştirilmiş; ek olarak kullanıcı arayüzüne dinamik genişletme "
    "özelliği ve ML skoru hesaplama iyileştirmesi eklenmiştir."
)

add_heading(doc, "7.2 Öneriler", size=13)
for o in [
    "Model zenginleştirme: Amfipatisite ve izoelektrik nokta gibi ek biyofiziksel özellikler eklenerek doğruluk artırılabilir.",
    "In vitro doğrulama: Sistemin en iyi olarak önerdiği adaylar gerçek laboratuvar ortamında test edilmelidir.",
    "Veritabanı genişletme: APD3, CAMP gibi ek veritabanları entegre edilerek eğitim seti zenginleştirilebilir.",
    "Aktif öğrenme: Laboratuvarda başarılı olan peptitler sisteme geri beslenerek model sürekli iyileştirilebilir.",
    "Ambalaj entegrasyonu: Bulunan AMP'lerin akıllı ambalaj filmlerine nasıl dahil edileceği gıda mühendisliği perspektifinden ayrıca araştırılmalıdır.",
    "Ölçekleme: Sistem, bulut altyapısında daha büyük peptit havuzlarını işleyecek şekilde ölçeklendirilebilir.",
]:
    add_bullet(doc, o)

# ═══════════════════════════════════════════════════════════════════════════
#  KAYNAKÇA
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "KAYNAKÇA", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
for ref in [
    "Dai, Y., et al. (2026). Utilizing Lactiplantibacillus plantarum to combat foodborne pathogens. Chemical Engineering Journal. https://doi.org/10.1016/j.cej.2026.173879",
    "Farrukh, S., et al. (2025). Antibiotic resistance and preventive strategies in foodborne pathogenic bacteria. Food Science and Biotechnology. https://doi.org/10.1007/s10068-024-01767-x",
    "Huan, Y., et al. (2020). Antimicrobial peptides: Classification, design, application and research progress. Frontiers in Microbiology, 11, 582779.",
    "Hu, X., et al. (2025). A food-grade antimicrobial peptide from Lactococcus lactis. Food Control. https://doi.org/10.1016/j.foodcont.2025.111922",
    "Li, J., et al. (2021). Plant antimicrobial peptides: Structures, functions, and applications. Botanical Studies, 62, 12.",
    "Oraç, H. (2024). Sustainable sources of bioactive peptides. Turkish Journal of Agriculture – Food Science and Technology, 12(5), 855–866.",
    "Rai, A., et al. (2016). Antimicrobial peptides as natural bio-preservative to enhance the shelf-life of food. Journal of Food Science and Technology, 53(6), 2278–2288.",
    "Song, L., et al. (2026). Sesquicilins A–H: 11-mer peptaibols with selective antibacterial activity. Aquaculture and Fisheries. https://doi.org/10.1016/j.aaf.2026.03.004",
    "Wang, H., et al. (2025). From precision synthesis to cross-industry applications: The future of emerging peptide technologies. Pharmacological Research, 107839.",
    "Xie, Q., et al. (2025). Advanced antimicrobial peptide-based biomaterials for food safety applications. LWT – Food Science and Technology, 188, 118884.",
    "Zhang, Y., et al. (2024). Emerging trends and challenges in antimicrobial peptide research. Processes, 12(5), 527.",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    run = p.add_run(ref)
    set_font(run, size=11)

# ═══════════════════════════════════════════════════════════════════════════
#  EKLER
# ═══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "EK A – Biyolojik Kural Yapılandırma Dosyası (biological_rules.json)", size=13)
add_body(doc,
    "Sistemin biyolojik filtreleme kuralları biological_rules.json dosyasında saklanmakta "
    "ve pipeline başlatılırken yüklenmektedir. Aşağıda dosyanın tam içeriği verilmiştir:"
)
p = doc.add_paragraph()
run = p.add_run(
    '{\n'
    '  "target_bacteria": "Listeria monocytogenes",\n'
    '  "length_range": [15, 30],\n'
    '  "charge_range": [3, 9],\n'
    '  "hydrophobicity_range": [0.4, 0.65],\n'
    '  "key_amino_acids": ["K", "R", "L", "I", "F"]\n'
    '}'
)
run.font.name = "Courier New"
run.font.size = Pt(10)

doc.add_page_break()
add_heading(doc, "EK B – CSV Yükleme Formatı", size=13)
add_body(doc,
    "Kullanıcının kendi peptit listesini yüklemek için kullanabileceği CSV formatı "
    "aşağıda gösterilmiştir. Dosyanın 'sequence' sütununu içermesi zorunludur:"
)
p = doc.add_paragraph()
run = p.add_run("sequence\nKWKLFKKIGAVLKVL\nLLKKLLKKLLKK\nRWKRLKRLKRLK")
run.font.name = "Courier New"
run.font.size = Pt(10)

# ── Kaydet ─────────────────────────────────────────────────────────────────
out = "BioPreserve_AI_Final_Raporu_Haziran2026.docx"
doc.save(out)
print(f"Rapor oluşturuldu: {out}")
